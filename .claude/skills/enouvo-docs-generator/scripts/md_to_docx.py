#!/usr/bin/env python3
"""Convert Markdown to DOCX with watermark, header, and footer support.

Converts markdown files to Word documents with optional watermark/logo,
header (logo + title), and footer (page numbers).

Features:
- Proper text formatting (bold, italic, code)
- Table rendering support
- Mermaid diagram pre-rendering (requires mmdc CLI)
- Watermark on all pages
- Header with logo and title
- Footer with page numbers
- Syntax highlighting for code blocks

Usage:
    python md_to_docx.py input.md [output.docx] [--watermark logo.png]
    python md_to_docx.py input.md --header-logo logo.png --watermark watermark.png

Requirements:
    pip install python-docx markdown pygments pymdown-extensions Pillow

Optional (for mermaid diagrams):
    npm install -g @mermaid-js/mermaid-cli
"""

import argparse
import io
import os
import re
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import List, Tuple, Optional

def install_package(package_name: str) -> bool:
    """Try to install a package using pip (handles PEP 668 externally-managed environments)."""
    # Try different pip install strategies for compatibility
    strategies = [
        [sys.executable, '-m', 'pip', 'install', '--user', package_name],
        [sys.executable, '-m', 'pip', 'install', '--user', '--break-system-packages', package_name],
    ]
    for cmd in strategies:
        try:
            subprocess.run(cmd, check=True, capture_output=True)
            return True
        except subprocess.CalledProcessError:
            continue
    print(f"Auto-install failed for {package_name}. Please install manually:")
    print(f"  pip install --user {package_name}")
    return False


def check_and_install(package_name: str, import_name: str) -> bool:
    """Check if package is importable, install if not."""
    try:
        __import__(import_name)
        return True
    except ImportError:
        print(f"Installing {package_name}...")
        if install_package(package_name):
            # Clear import caches and retry
            import importlib
            importlib.invalidate_caches()
            try:
                __import__(import_name)
                return True
            except ImportError:
                pass
        print(f"Failed to install {package_name}. Please install manually:")
        print(f"  pip install {package_name}")
        return False


# Auto-install core dependencies
required_packages = [
    ('python-docx', 'docx'),
    ('markdown', 'markdown'),
    ('Pillow', 'PIL'),
]

for pkg, imp in required_packages:
    if not check_and_install(pkg, imp):
        sys.exit(1)

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from PIL import Image
import markdown
from markdown.extensions.tables import TableExtension

# Auto-install and check optional dependencies
import importlib

PYGMENTS_AVAILABLE = False
try:
    from pygments import highlight
    from pygments.lexers import get_lexer_by_name, guess_lexer
    from pygments.token import Token
    PYGMENTS_AVAILABLE = True
except ImportError:
    print("Installing Pygments for syntax highlighting...")
    if install_package('Pygments'):
        importlib.invalidate_caches()
        try:
            from pygments import highlight
            from pygments.lexers import get_lexer_by_name, guess_lexer
            from pygments.token import Token
            PYGMENTS_AVAILABLE = True
        except ImportError:
            print("Warning: Pygments installed but import failed")

PYMDOWNX_AVAILABLE = False
try:
    import pymdownx
    PYMDOWNX_AVAILABLE = True
except ImportError:
    print("Installing pymdown-extensions for better code blocks...")
    if install_package('pymdown-extensions'):
        importlib.invalidate_caches()
        try:
            import pymdownx
            PYMDOWNX_AVAILABLE = True
        except ImportError:
            print("Warning: pymdown-extensions installed but import failed")

# Enouvo brand colors
ENOUVO_RED = RGBColor(0xE7, 0x4C, 0x3C)  # #e74c3c
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY = RGBColor(0x33, 0x33, 0x33)  # Darker for better visibility
CODE_BG = RGBColor(0xF8, 0xF8, 0xF8)

# Syntax highlighting colors (GitHub-style)
SYNTAX_COLORS = {
    Token.Keyword: RGBColor(0xD7, 0x3A, 0x49),
    Token.Keyword.Declaration: RGBColor(0xD7, 0x3A, 0x49),
    Token.Keyword.Namespace: RGBColor(0xD7, 0x3A, 0x49),
    Token.Keyword.Type: RGBColor(0xD7, 0x3A, 0x49),
    Token.Name.Function: RGBColor(0x6F, 0x42, 0xC1),
    Token.Name.Class: RGBColor(0x6F, 0x42, 0xC1),
    Token.Name.Builtin: RGBColor(0x00, 0x5C, 0xC5),
    Token.String: RGBColor(0x03, 0x2F, 0x62),
    Token.String.Single: RGBColor(0x03, 0x2F, 0x62),
    Token.String.Double: RGBColor(0x03, 0x2F, 0x62),
    Token.Number: RGBColor(0x00, 0x5C, 0xC5),
    Token.Number.Integer: RGBColor(0x00, 0x5C, 0xC5),
    Token.Number.Float: RGBColor(0x00, 0x5C, 0xC5),
    Token.Comment: RGBColor(0x6A, 0x73, 0x7D),
    Token.Comment.Single: RGBColor(0x6A, 0x73, 0x7D),
    Token.Comment.Multiline: RGBColor(0x6A, 0x73, 0x7D),
    Token.Operator: RGBColor(0xD7, 0x3A, 0x49),
    Token.Name.Tag: RGBColor(0x22, 0x86, 0x3A),
    Token.Name.Attribute: RGBColor(0x6F, 0x42, 0xC1),
}


def install_dependency(package: str) -> bool:
    """Attempt to install a Python package via pip."""
    print(f"{package} not found. Attempting to install...")
    try:
        result = subprocess.run(
            [sys.executable, '-m', 'pip', 'install', package, '--user', '-q'],
            capture_output=True, timeout=60
        )
        if result.returncode == 0:
            print(f"{package} installed successfully!")
            return True
        result = subprocess.run(
            [sys.executable, '-m', 'pip', 'install', package, '-q'],
            capture_output=True, timeout=60
        )
        return result.returncode == 0
    except Exception:
        return False


def ensure_pygments() -> bool:
    """Ensure Pygments is available."""
    global PYGMENTS_AVAILABLE
    if PYGMENTS_AVAILABLE:
        return True
    if install_dependency('pygments'):
        try:
            from pygments import highlight
            PYGMENTS_AVAILABLE = True
            return True
        except ImportError:
            pass
    return False


def ensure_pymdownx() -> bool:
    """Ensure pymdown-extensions is available."""
    global PYMDOWNX_AVAILABLE
    if PYMDOWNX_AVAILABLE:
        return True
    if install_dependency('pymdown-extensions'):
        try:
            import pymdownx
            PYMDOWNX_AVAILABLE = True
            return True
        except ImportError:
            pass
    return False


def check_mermaid_cli() -> bool:
    """Check if mermaid-cli (mmdc) is installed."""
    try:
        result = subprocess.run(['mmdc', '--version'], capture_output=True, timeout=5)
        return result.returncode == 0
    except (subprocess.SubprocessError, FileNotFoundError):
        return False


def install_mermaid_cli() -> bool:
    """Attempt to install mermaid-cli via npm."""
    print("Mermaid CLI not found. Attempting to install...")
    try:
        npm_check = subprocess.run(['npm', '--version'], capture_output=True, timeout=10)
        if npm_check.returncode != 0:
            print("npm not found. Cannot auto-install mermaid-cli.")
            return False
    except (subprocess.SubprocessError, FileNotFoundError):
        print("npm not found. Cannot auto-install mermaid-cli.")
        return False
    try:
        print("Running: npm install -g @mermaid-js/mermaid-cli")
        result = subprocess.run(
            ['npm', 'install', '-g', '@mermaid-js/mermaid-cli'],
            capture_output=True,
            timeout=120
        )
        if result.returncode == 0:
            print("Mermaid CLI installed successfully!")
            return True
        stderr = result.stderr.decode('utf-8', errors='ignore')
        print(f"Installation failed: {stderr}")
        return False
    except subprocess.TimeoutExpired:
        print("Installation timed out.")
        return False
    except (subprocess.SubprocessError, FileNotFoundError) as e:
        print(f"Installation error: {e}")
        return False


def ensure_mermaid_cli() -> bool:
    """Ensure mermaid-cli is available, installing if necessary."""
    if check_mermaid_cli():
        return True
    if install_mermaid_cli():
        return check_mermaid_cli()
    return False


def render_mermaid_diagram(mermaid_code: str, output_path: str) -> bool:
    """Render mermaid diagram to high-resolution PNG (5x scale)."""
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as f:
            f.write(mermaid_code)
            input_path = f.name
        # 5x scale = very high resolution source for crisp downscaling in output
        result = subprocess.run(
            ['mmdc', '-i', input_path, '-o', output_path, '-b', 'transparent', '-s', '5'],
            capture_output=True, timeout=60
        )
        os.unlink(input_path)
        return result.returncode == 0
    except (subprocess.SubprocessError, FileNotFoundError, OSError):
        return False


def extract_title_from_markdown(md_content: str) -> str:
    """Extract the first H1 title from markdown content."""
    h1_pattern = r'^#\s+(.+?)$|^(.+?)\n=+$'
    match = re.search(h1_pattern, md_content, re.MULTILINE)
    if match:
        return (match.group(1) or match.group(2)).strip()
    return 'Document'


def create_element(tag: str) -> OxmlElement:
    """Create an OxmlElement with the given tag."""
    return OxmlElement(tag)


def set_cell_shading(cell, color: str):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number(paragraph):
    """Add page number field to paragraph."""
    run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = create_element('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    paragraph.add_run(" of ")

    run2 = paragraph.add_run()
    fldChar3 = create_element('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')

    instrText2 = create_element('w:instrText')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    run2._r.append(fldChar3)
    run2._r.append(instrText2)
    run2._r.append(fldChar4)


def convert_image_for_docx(image_path: str) -> io.BytesIO:
    """Convert image to PNG format in memory for python-docx compatibility.

    python-docx doesn't support webp format directly, so we convert
    any image to PNG in memory before adding to the document.
    Preserves original colors without fading.
    """
    img = Image.open(image_path)

    # Convert palette mode to RGBA to preserve colors
    if img.mode == 'P':
        img = img.convert('RGBA')

    # Keep RGBA images as-is (preserves transparency and color intensity)
    # Only convert other modes to RGB
    if img.mode not in ('RGB', 'RGBA'):
        img = img.convert('RGB')

    # Save to bytes buffer as PNG (supports both RGB and RGBA)
    buffer = io.BytesIO()
    img.save(buffer, format='PNG')
    buffer.seek(0)
    return buffer


def setup_document_styles(doc: Document):
    """Set up custom styles for the document."""
    styles = doc.styles

    # Heading 1
    if 'Heading 1' in styles:
        h1 = styles['Heading 1']
        h1.font.size = Pt(24)
        h1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        h1.font.bold = True

    # Heading 2
    if 'Heading 2' in styles:
        h2 = styles['Heading 2']
        h2.font.size = Pt(18)
        h2.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        h2.font.bold = True

    # Heading 3
    if 'Heading 3' in styles:
        h3 = styles['Heading 3']
        h3.font.size = Pt(14)
        h3.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
        h3.font.bold = True

    # Normal text
    normal = styles['Normal']
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK_GRAY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_header_with_logo(doc: Document, title: str, logo_path: str = None):
    """Add compact header with logo and title (enterprise style)."""
    section = doc.sections[0]
    header = section.header

    # Create a table for header layout (logo left, title right)
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    table.columns[0].width = Inches(1.2)
    table.columns[1].width = Inches(5.3)

    # Logo cell - visible
    logo_cell = table.cell(0, 0)
    logo_para = logo_cell.paragraphs[0]
    logo_para.paragraph_format.space_after = Pt(0)
    if logo_path and Path(logo_path).exists():
        try:
            run = logo_para.add_run()
            img_buffer = convert_image_for_docx(logo_path)
            run.add_picture(img_buffer, height=Inches(0.4))  # Clear logo
        except Exception as e:
            print(f"Warning: Could not load header logo: {e}")

    # Title cell - readable font
    title_cell = table.cell(0, 1)
    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_para.paragraph_format.space_after = Pt(0)
    run = title_para.add_run(title)
    run.font.size = Pt(11)  # Readable size
    run.font.color.rgb = LIGHT_GRAY

    # Add red separator line
    border_para = header.add_paragraph()
    border_para.paragraph_format.space_before = Pt(2)
    border_para.paragraph_format.space_after = Pt(12)  # Space to content
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')  # Visible line
    bottom.set(qn('w:color'), 'E74C3C')
    pBdr.append(bottom)
    border_para._p.get_or_add_pPr().append(pBdr)


def add_footer_with_page_numbers(doc: Document):
    """Add compact footer with red separator line and page numbers."""
    section = doc.sections[0]
    footer = section.footer

    # Add red separator line at top of footer - minimal spacing
    border_para = footer.add_paragraph()
    border_para.paragraph_format.space_before = Pt(4)  # Minimal space from content
    border_para.paragraph_format.space_after = Pt(2)
    pBdr = OxmlElement('w:pBdr')
    top_border = OxmlElement('w:top')
    top_border.set(qn('w:val'), 'single')
    top_border.set(qn('w:sz'), '8')  # Same as header
    top_border.set(qn('w:color'), 'E74C3C')
    pBdr.append(top_border)
    border_para._p.get_or_add_pPr().append(pBdr)

    # Page numbers
    para = footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.add_run("Page ")
    add_page_number(para)
    for run in para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = LIGHT_GRAY


def add_watermark(doc: Document, image_path: str, opacity: float = 0.12):
    """Add centered watermark image behind text on all pages using DrawingML anchor."""
    if not Path(image_path).exists():
        return

    section = doc.sections[0]
    header = section.header

    # Get image dimensions and calculate size (60% of page width for visibility)
    img = Image.open(image_path)
    img_width, img_height = img.size
    aspect_ratio = img_height / img_width

    page_width_emu = int(section.page_width)
    page_height_emu = int(section.page_height)
    wm_width_emu = int(page_width_emu * 0.6)  # 60% of page width
    wm_height_emu = int(wm_width_emu * aspect_ratio)

    try:
        # Create semi-transparent version
        img = img.convert('RGBA')
        alpha = img.split()[3]
        alpha = alpha.point(lambda p: int(p * opacity))
        img.putalpha(alpha)

        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            img.save(tmp.name, 'PNG')

            # Add paragraph to header for watermark
            wm_para = header.add_paragraph()
            run = wm_para.add_run()

            # Add picture and get the inline element
            picture = run.add_picture(tmp.name, width=Emu(wm_width_emu))
            inline = picture._inline

            # Find the drawing element (parent of inline)
            drawing = inline.getparent()

            # Convert inline to anchor for precise positioning
            anchor = _create_anchor_from_inline(inline, wm_width_emu, wm_height_emu, page_width_emu, page_height_emu)

            # Replace inline with anchor inside the drawing
            if drawing is not None:
                drawing.remove(inline)
                drawing.append(anchor)

            # Hide the paragraph (zero height)
            wm_para.paragraph_format.space_before = Pt(0)
            wm_para.paragraph_format.space_after = Pt(0)
            wm_para.paragraph_format.line_spacing = Pt(0)

            os.unlink(tmp.name)
    except Exception as e:
        print(f"Warning: Could not add watermark: {e}")


def _create_anchor_from_inline(inline, width_emu: int, height_emu: int, page_width_emu: int, page_height_emu: int):
    """Create a DrawingML anchor element for centered page watermark."""
    from lxml import etree

    # Calculate true center position on the page
    h_offset = (page_width_emu - width_emu) // 2
    v_offset = (page_height_emu - height_emu) // 2

    # Namespaces
    WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'

    # Create anchor element using lxml directly
    anchor = etree.Element(f'{{{WP_NS}}}anchor')
    anchor.set('distT', '0')
    anchor.set('distB', '0')
    anchor.set('distL', '0')
    anchor.set('distR', '0')
    anchor.set('simplePos', '0')
    anchor.set('relativeHeight', '0')
    anchor.set('behindDoc', '1')  # Behind document text
    anchor.set('locked', '0')
    anchor.set('layoutInCell', '1')
    anchor.set('allowOverlap', '1')

    # Simple position (not used but required)
    simple_pos = etree.SubElement(anchor, f'{{{WP_NS}}}simplePos')
    simple_pos.set('x', '0')
    simple_pos.set('y', '0')

    # Horizontal position - center of page
    pos_h = etree.SubElement(anchor, f'{{{WP_NS}}}positionH')
    pos_h.set('relativeFrom', 'page')
    pos_offset_h = etree.SubElement(pos_h, f'{{{WP_NS}}}posOffset')
    pos_offset_h.text = str(h_offset)

    # Vertical position - center of page
    pos_v = etree.SubElement(anchor, f'{{{WP_NS}}}positionV')
    pos_v.set('relativeFrom', 'page')
    pos_offset_v = etree.SubElement(pos_v, f'{{{WP_NS}}}posOffset')
    pos_offset_v.text = str(v_offset)

    # Extent (size)
    extent = etree.SubElement(anchor, f'{{{WP_NS}}}extent')
    extent.set('cx', str(width_emu))
    extent.set('cy', str(height_emu))

    # Effect extent
    effect_extent = etree.SubElement(anchor, f'{{{WP_NS}}}effectExtent')
    effect_extent.set('l', '0')
    effect_extent.set('t', '0')
    effect_extent.set('r', '0')
    effect_extent.set('b', '0')

    # Wrap none (no text wrapping - image is behind)
    etree.SubElement(anchor, f'{{{WP_NS}}}wrapNone')

    # Doc properties
    doc_pr = etree.SubElement(anchor, f'{{{WP_NS}}}docPr')
    doc_pr.set('id', '1')
    doc_pr.set('name', 'Watermark')

    # Copy the graphic element from the inline
    graphic = inline.find(f'{{{A_NS}}}graphic')
    if graphic is not None:
        import copy
        anchor.append(copy.deepcopy(graphic))

    return anchor


class MarkdownToDocxConverter:
    """Convert markdown content to DOCX document."""

    def __init__(self, doc: Document, temp_dir: Path):
        self.doc = doc
        self.temp_dir = temp_dir
        self.has_mermaid = ensure_mermaid_cli()
        self.current_list_level = 0
        self.in_code_block = False
        self.code_language = None
        self.code_content = []

    def add_heading(self, text: str, level: int):
        """Add a heading to the document."""
        para = self.doc.add_heading(text, level=min(level, 9))
        if level == 1:
            # Add red underline for H1
            para.paragraph_format.space_after = Pt(6)

    def add_paragraph(self, text: str):
        """Add a paragraph with inline formatting."""
        para = self.doc.add_paragraph()
        self._add_formatted_text(para, text)
        return para

    def _add_formatted_text(self, para, text: str):
        """Parse and add text with inline formatting (bold, italic, code, links)."""
        # Pattern for inline formatting
        patterns = [
            (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),
            (r'\*\*(.+?)\*\*', 'bold'),
            (r'__(.+?)__', 'bold'),
            (r'\*(.+?)\*', 'italic'),
            (r'_(.+?)_', 'italic'),
            (r'`(.+?)`', 'code'),
            (r'\[(.+?)\]\((.+?)\)', 'link'),
        ]

        # Simple approach: process text segments
        remaining = text
        while remaining:
            earliest_match = None
            earliest_pos = len(remaining)
            match_type = None

            for pattern, fmt_type in patterns:
                match = re.search(pattern, remaining)
                if match and match.start() < earliest_pos:
                    earliest_match = match
                    earliest_pos = match.start()
                    match_type = fmt_type

            if earliest_match:
                # Add text before match
                if earliest_pos > 0:
                    para.add_run(remaining[:earliest_pos])

                # Add formatted text
                if match_type == 'bold_italic':
                    run = para.add_run(earliest_match.group(1))
                    run.bold = True
                    run.italic = True
                elif match_type == 'bold':
                    run = para.add_run(earliest_match.group(1))
                    run.bold = True
                elif match_type == 'italic':
                    run = para.add_run(earliest_match.group(1))
                    run.italic = True
                elif match_type == 'code':
                    run = para.add_run(earliest_match.group(1))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                elif match_type == 'link':
                    # Links shown as text (DOCX hyperlinks are complex)
                    run = para.add_run(earliest_match.group(1))
                    run.font.color.rgb = ENOUVO_RED
                    run.underline = True

                remaining = remaining[earliest_match.end():]
            else:
                para.add_run(remaining)
                break

    def add_code_block(self, code: str, language: str = None):
        """Add a code block with syntax highlighting."""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

        # Add subtle left border + very light background for code readability
        pPr = para._p.get_or_add_pPr()

        # Left border
        pBdr = OxmlElement('w:pBdr')
        left_border = OxmlElement('w:left')
        left_border.set(qn('w:val'), 'single')
        left_border.set(qn('w:sz'), '18')
        left_border.set(qn('w:color'), 'DDDDDD')
        pBdr.append(left_border)
        pPr.append(pBdr)

        # Very light background (semi-transparent effect - watermark visible through)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), 'FAFAFA')  # Nearly white - watermark shows through
        pPr.append(shd)

        if PYGMENTS_AVAILABLE and language:
            self._add_highlighted_code(para, code, language)
        else:
            run = para.add_run(code)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

    def _add_highlighted_code(self, para, code: str, language: str):
        """Add syntax-highlighted code."""
        try:
            lexer = get_lexer_by_name(language)
        except Exception:
            try:
                lexer = guess_lexer(code)
            except Exception:
                run = para.add_run(code)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                return

        from pygments import lex
        tokens = list(lex(code, lexer))

        for token_type, token_value in tokens:
            run = para.add_run(token_value)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

            # Apply color based on token type
            color = None
            for t_type in [token_type] + list(token_type.split()):
                if t_type in SYNTAX_COLORS:
                    color = SYNTAX_COLORS[t_type]
                    break

            if color:
                run.font.color.rgb = color

            # Make keywords bold
            if token_type in Token.Keyword:
                run.bold = True

            # Make comments italic
            if token_type in Token.Comment:
                run.italic = True

    def add_table(self, rows: List[List[str]], has_header: bool = True):
        """Add a table to the document."""
        if not rows:
            return

        num_cols = len(rows[0])
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        table.autofit = True

        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_text.strip()

                # Style header row
                if i == 0 and has_header:
                    set_cell_shading(cell, 'E74C3C')
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.bold = True
                elif i % 2 == 0:
                    set_cell_shading(cell, 'FAFAFA')  # Lighter for watermark visibility

    def add_list_item(self, text: str, ordered: bool = False, level: int = 0):
        """Add a list item."""
        para = self.doc.add_paragraph(style='List Bullet' if not ordered else 'List Number')
        para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        self._add_formatted_text(para, text)

    def add_image(self, image_path: str, alt_text: str = ''):
        """Add an image to the document."""
        if not Path(image_path).exists():
            return

        try:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(image_path, width=Inches(5))
        except Exception as e:
            print(f"Warning: Could not add image: {e}")

    def add_mermaid_diagram(self, mermaid_code: str):
        """Render and add a mermaid diagram."""
        if not self.has_mermaid:
            # Fallback: show code
            self.add_code_block(mermaid_code, 'text')
            return

        diagram_id = hash(mermaid_code) & 0xFFFFFFFF
        png_path = self.temp_dir / f"mermaid_{diagram_id}.png"

        if render_mermaid_diagram(mermaid_code, str(png_path)):
            self.add_image(str(png_path), 'Mermaid Diagram')
        else:
            self.add_code_block(mermaid_code, 'text')

    def add_blockquote(self, text: str):
        """Add a blockquote."""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)

        # Add left border via shading
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')
        left.set(qn('w:color'), 'E74C3C')
        pBdr.append(left)
        pPr.append(pBdr)

        run = para.add_run(text)
        run.italic = True

    def add_horizontal_rule(self):
        """Add a horizontal rule."""
        para = self.doc.add_paragraph()
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), 'DDDDDD')
        pBdr.append(bottom)
        pPr.append(pBdr)


def parse_and_convert(md_content: str, converter: MarkdownToDocxConverter):
    """Parse markdown and convert to DOCX using the converter."""
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_language = None
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        code_match = re.match(r'^(\s*)```(\w*)\s*$', line)
        if code_match:
            if not in_code_block:
                in_code_block = True
                code_language = code_match.group(2) or None
                code_lines = []
            else:
                # End of code block
                code_content = '\n'.join(code_lines)
                if code_language == 'mermaid':
                    converter.add_mermaid_diagram(code_content)
                else:
                    converter.add_code_block(code_content, code_language)
                in_code_block = False
                code_language = None
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []

            # Skip separator row
            if re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
                i += 1
                continue

            # Parse table row
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table
            converter.add_table(table_rows)
            in_table = False
            table_rows = []

        # Headings
        h_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if h_match:
            level = len(h_match.group(1))
            converter.add_heading(h_match.group(2), level)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^(\*{3,}|-{3,}|_{3,})\s*$', line):
            converter.add_horizontal_rule()
            i += 1
            continue

        # Blockquote
        bq_match = re.match(r'^>\s*(.*)$', line)
        if bq_match:
            converter.add_blockquote(bq_match.group(1))
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ol_match:
            indent = len(ol_match.group(1)) // 2
            converter.add_list_item(ol_match.group(2), ordered=True, level=indent)
            i += 1
            continue

        # Unordered list
        ul_match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
        if ul_match:
            indent = len(ul_match.group(1)) // 2
            converter.add_list_item(ul_match.group(2), ordered=False, level=indent)
            i += 1
            continue

        # Image
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', line.strip())
        if img_match:
            converter.add_image(img_match.group(2), img_match.group(1))
            i += 1
            continue

        # Regular paragraph (non-empty lines)
        if line.strip():
            converter.add_paragraph(line)

        i += 1

    # Handle any remaining table
    if in_table and table_rows:
        converter.add_table(table_rows)


def md_to_docx(
    input_path: str,
    output_path: str = None,
    watermark_path: str = None,
    header_logo_path: str = None,
    custom_title: str = None,
    show_footer: bool = True
) -> str:
    """Convert markdown file to DOCX with optional watermark, header, and footer.

    Args:
        input_path: Path to the markdown file
        output_path: Output DOCX path (defaults to input_path with .docx extension)
        watermark_path: Optional path to watermark image
        header_logo_path: Optional path to header logo image
        custom_title: Optional custom title (defaults to first H1)
        show_footer: Whether to show page numbers in footer

    Returns:
        Path to the created DOCX file
    """
    input_path = Path(input_path)

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    if output_path is None:
        output_path = input_path.with_suffix('.docx')
    output_path = Path(output_path)

    # Read markdown content
    md_content = input_path.read_text(encoding='utf-8')

    # Extract title
    title = custom_title or extract_title_from_markdown(md_content)

    # Create document
    doc = Document()
    setup_document_styles(doc)

    # Add watermark first (subtle branding in header area)
    if watermark_path:
        add_watermark(doc, watermark_path)

    # Add header with logo
    if header_logo_path:
        add_header_with_logo(doc, title, header_logo_path)

    # Add footer
    if show_footer:
        add_footer_with_page_numbers(doc)

    # Convert markdown to DOCX
    with tempfile.TemporaryDirectory() as temp_dir:
        converter = MarkdownToDocxConverter(doc, Path(temp_dir))
        parse_and_convert(md_content, converter)

    # Save document
    doc.save(str(output_path))

    return str(output_path)


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown to DOCX with optional watermark, header, and footer',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    %(prog)s document.md
    %(prog)s document.md output.docx
    %(prog)s document.md --watermark watermark.png
    %(prog)s document.md --header-logo logo.png
    %(prog)s document.md --header-logo logo.webp --watermark watermark.png
    %(prog)s document.md --header-logo logo.png --title "My Document"
    %(prog)s document.md --no-footer

Features:
    - Text formatting (bold, italic, code)
    - Table rendering with styled headers
    - Mermaid diagram support (requires: npm install -g @mermaid-js/mermaid-cli)
    - Syntax highlighting for code blocks
    - Watermark on all pages
    - Header with logo and document title
    - Footer with page numbers
        """
    )

    parser.add_argument('input', help='Input markdown file path')
    parser.add_argument('output', nargs='?', help='Output DOCX file path (default: input.docx)')
    parser.add_argument('--watermark', '-w', help='Path to watermark image')
    parser.add_argument('--header-logo', '-l', help='Path to header logo image')
    parser.add_argument('--title', '-t', help='Document title for header')
    parser.add_argument('--no-footer', action='store_true', help='Disable page numbers')

    args = parser.parse_args()

    # Check dependencies
    if check_mermaid_cli():
        print("Mermaid CLI detected - diagrams will be rendered")
    else:
        print("Note: Mermaid CLI not found - diagrams will show as code")

    if ensure_pygments():
        print("Pygments detected - code syntax highlighting enabled")
    else:
        print("Warning: Pygments not available - no syntax highlighting")

    try:
        output = md_to_docx(
            input_path=args.input,
            output_path=args.output,
            watermark_path=args.watermark,
            header_logo_path=args.header_logo,
            custom_title=args.title,
            show_footer=not args.no_footer
        )
        print(f"DOCX created: {output}")
    except FileNotFoundError as e:
        print(f"Error: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"Error converting file: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
