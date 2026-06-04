#!/usr/bin/env python3
"""Extract text from PDF or DOCX files for spec-forge ingestion."""

import sys
import hashlib
from pathlib import Path


def extract_pdf(path: Path) -> str:
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[Page {i}]\n{text}")
        return "\n\n".join(pages)
    except ImportError:
        pass

    try:
        import pdfminer.high_level
        return pdfminer.high_level.extract_text(str(path))
    except ImportError:
        pass

    raise RuntimeError(
        "No PDF library available. Install one: pip install pdfplumber  OR  pip install pdfminer.six"
    )


def extract_docx(path: Path) -> str:
    try:
        import docx
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = docx.Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def file_hash(path: Path) -> str:
    h = hashlib.md5()
    h.update(path.read_bytes())
    return h.hexdigest()


def main():
    if len(sys.argv) < 2:
        print("Usage: extract_text.py <file> [--hash]", file=sys.stderr)
        sys.exit(1)

    path = Path(sys.argv[1])
    if not path.exists():
        print(f"File not found: {path}", file=sys.stderr)
        sys.exit(1)

    if "--hash" in sys.argv:
        print(file_hash(path))
        return

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = extract_pdf(path)
    elif suffix in (".docx", ".doc"):
        text = extract_docx(path)
    else:
        text = path.read_text(encoding="utf-8", errors="replace")

    print(text)


if __name__ == "__main__":
    main()
